#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 consumer-evtol-design-2026.md 导出为 Word (.docx) 与 PDF (.pdf)。

依赖：pypandoc_binary、python-docx、docx2pdf（Windows 上 PDF 需本机安装 Microsoft Word）

用法：
    python export_design_doc.py
    python export_design_doc.py --md path/to/file.md --out-dir ./output
"""

from __future__ import annotations

import argparse
import shutil
import sys
from datetime import date
from pathlib import Path

try:
    import pypandoc
except ImportError:
    print("请先安装：pip install pypandoc_binary python-docx docx2pdf", file=sys.stderr)
    sys.exit(1)

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

# 中文字体（Word 内置）
FONT_BODY = "微软雅黑"
FONT_CODE = "Consolas"
FONT_TITLE = "微软雅黑"


def ensure_pandoc() -> None:
    """确保 pandoc 可执行文件可用（pypandoc_binary 自带）。"""
    try:
        pypandoc.get_pandoc_version()
    except OSError:
        print("pandoc 不可用，请执行：pip install pypandoc_binary", file=sys.stderr)
        sys.exit(1)


def build_reference_docx(ref_path: Path) -> None:
    """生成 pandoc 参考模板：标题层级、正文、表格基础样式。"""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    def _style_font(style_name: str, size_pt: float, bold: bool = False, color: RGBColor | None = None):
        st = doc.styles[style_name]
        st.font.name = FONT_BODY
        st.font.size = Pt(size_pt)
        st.font.bold = bold
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        if color:
            st.font.color.rgb = color
        pf = st.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.25
        pf.space_after = Pt(6)

    _style_font("Normal", 10.5)
    _style_font("Heading 1", 18, bold=True, color=RGBColor(0x1A, 0x47, 0x8C))
    _style_font("Heading 2", 14, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    _style_font("Heading 3", 12, bold=True, color=RGBColor(0x40, 0x40, 0x40))
    _style_font("Heading 4", 11, bold=True)

    # 占位段落，pandoc 会按样式名映射
    doc.add_heading("Heading 1", level=1)
    doc.add_heading("Heading 2", level=2)
    doc.add_heading("Heading 3", level=3)
    doc.add_paragraph("正文示例 Normal。")

    ref_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ref_path))


def _set_run_font(run, name: str, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def polish_docx(docx_path: Path) -> None:
    """导出后统一中文字体、表格与代码块样式。"""
    doc = Document(str(docx_path))

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title", "Subtitle"):
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = FONT_BODY
            st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        is_code = "Source Code" in style or "Verbatim" in style or style == "Preformatted"
        for run in para.runs:
            if is_code:
                _set_run_font(run, FONT_CODE, 9.0)
            elif style.startswith("Heading"):
                _set_run_font(run, FONT_TITLE, run.font.size.pt if run.font.size else None, True)
            else:
                _set_run_font(run, FONT_BODY, run.font.size.pt if run.font.size else 10.5)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_run_font(run, FONT_BODY, 9.5)

    doc.save(str(docx_path))


def prepare_md_with_metadata(md_path: Path, work_dir: Path) -> Path:
    """复制 MD 并加上 YAML 元数据（标题、目录）供 pandoc 使用。"""
    raw = md_path.read_text(encoding="utf-8")
    # 避免与 YAML title 重复：去掉 MD 首行一级标题
    lines = raw.splitlines()
    if lines and lines[0].startswith("# "):
        raw = "\n".join(lines[1:]).lstrip("\n")
    yaml = f"""---
title: "ZZZ eVTOL 飞控系列产品与 IO 规划"
subtitle: "ZZZ-EVTOL-H753 / F405 / H743 · 硬件设计规格书（2026）"
author: "ZZZ EVTOL FC Series"
date: "{date.today().isoformat()}"
lang: zh-CN
toc-title: "目录"
---

"""
    out = work_dir / f"{md_path.stem}_pandoc.md"
    out.write_text(yaml + raw, encoding="utf-8")
    return out


def md_to_docx(md_path: Path, docx_path: Path, ref_docx: Path) -> None:
    """Markdown → Word（pandoc + 参考模板）。"""
    extra_args = [
        f"--reference-doc={ref_docx}",
        "--toc",
        "--toc-depth=3",
        "--standalone",
    ]
    pypandoc.convert_file(
        str(md_path),
        "docx",
        outputfile=str(docx_path),
        extra_args=extra_args,
    )
    polish_docx(docx_path)


def docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Word → PDF（Windows：docx2pdf 调用 MS Word）。"""
    try:
        from docx2pdf import convert
    except ImportError:
        raise RuntimeError("请安装 docx2pdf：pip install docx2pdf") from None

    convert(str(docx_path), str(pdf_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="导出 eVTOL 设计 Markdown 为 Word/PDF")
    parser.add_argument(
        "--md",
        type=Path,
        default=Path(__file__).parent / "consumer-evtol-design-2026.md",
        help="源 Markdown 文件",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=Path(__file__).parent / "export",
        help="输出目录",
    )
    args = parser.parse_args()

    md_path: Path = args.md.resolve()
    out_dir: Path = args.out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    if not md_path.is_file():
        print(f"找不到 Markdown：{md_path}", file=sys.stderr)
        return 1

    stem = md_path.stem
    docx_path = out_dir / f"{stem}.docx"
    pdf_path = out_dir / f"{stem}.pdf"
    ref_docx = out_dir / "_reference_template.docx"

    ensure_pandoc()
    print(f"[1/4] 生成 Word 参考模板 …")
    build_reference_docx(ref_docx)

    print(f"[2/4] 准备 Markdown（YAML 元数据 + 目录）…")
    pandoc_md = prepare_md_with_metadata(md_path, out_dir)

    print(f"[3/4] Pandoc 转换 Markdown → Word …")
    print(f"      源：{md_path}")
    md_to_docx(pandoc_md, docx_path, ref_docx)
    print(f"      → {docx_path}")

    print(f"[4/4] 导出 PDF（需本机 Microsoft Word）…")
    try:
        docx_to_pdf(docx_path, pdf_path)
        print(f"      → {pdf_path}")
    except Exception as exc:
        print(f"      PDF 导出失败：{exc}", file=sys.stderr)
        print("      提示：请安装 Microsoft Word，或手动在 Word 中打开 docx 另存为 PDF。", file=sys.stderr)
        return 2

    # 清理临时模板（可选保留供调试）
    if ref_docx.exists():
        ref_docx.unlink()
    if pandoc_md.exists():
        pandoc_md.unlink()

    print(f"完成。")
    print(f"Word : {docx_path}")
    print(f"PDF  : {pdf_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
